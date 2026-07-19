import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
FIGURES = ROOT / "figures"
RESULTS = ROOT / "results"
CITE_NUMBERS = {
    key: i + 1 for i, key in enumerate([
        "breiman1984cart", "deng2012rrf", "deng2013grrf", "strobl2008conditional",
        "nicodemus2010correlation", "genuer2010selection", "louppe2013importance",
        "dormann2013collinearity", "breiman1996bagging", "li2002instability",
        "hastie2009esl", "rudin2019stop", "pedregosa2011sklearn", "uci2024",
        "demsar2006comparisons",
    ])
}
REF_NUMBERS = {
    "fig:synthetic": "1", "fig:sensitivity": "2", "fig:redundancy": "3",
    "fig:sweep": "4", "fig:trees": "5", "tab:stability": "6",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, value=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge in ("top", "start", "bottom", "end"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(value if edge in ("top", "bottom") else 120))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)


def add_table(doc, caption, frame, widths=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    table = doc.add_table(rows=1, cols=len(frame.columns))
    table.autofit = False
    table.style = "Table Grid"
    for j, column in enumerate(frame.columns):
        cell = table.rows[0].cells[j]
        cell.text = str(column)
        set_cell_shading(cell, "E8EEF5")
    for values in frame.itertuples(index=False):
        cells = table.add_row().cells
        for j, value in enumerate(values):
            cells[j].text = f"{value:.3f}" if isinstance(value, float) else str(value)
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            set_cell_margins(cell)
            if widths:
                cell.width = Inches(widths[j])
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(7.5)
                    if row is table.rows[0]:
                        run.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def clean_latex(text):
    text = re.sub(r"\\(?:begin|end)\{multicols\}\{?2?\}?", "", text)
    text = text.replace("\\begingroup", "").replace("\\endgroup", "").replace("\\let\\small\\scriptsize", "")
    text = re.sub(r"\\cite\{([^}]+)\}", lambda m: "[" + ", ".join(str(CITE_NUMBERS.get(k, k)) for k in m.group(1).split(",")) + "]", text)
    text = text.replace("\\%", "%").replace("\\&", "&")
    text = re.sub(r"\\texttt\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\emph\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\textbf\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\text\{([^}]*)\}", r"\1", text)
    text = re.sub(r"\\ref\{([^}]*)\}", lambda m: REF_NUMBERS.get(m.group(1), m.group(1)), text)
    text = text.replace("\\sqrt{1-\\rho^2}\\epsilon", "sqrt(1-rho^2) epsilon")
    text = text.replace("``", "“").replace("''", "”")
    text = text.replace("$", "")
    text = text.replace("Cram{\\'e}r", "Cramer's").replace("Carr{\\'e}", "Carre")
    text = text.replace("Dem{\\v{s}}ar", "Demsar").replace("Ga{\\\"e}l", "Gael")
    text = text.replace("\\alpha", "α").replace("\\rho", "ρ").replace("\\Delta", "Δ").replace("\\leq", "≤")
    text = text.replace("X_j", "Xⱼ").replace("X_p", "Xₚ").replace("\\in", " ∈ ")
    text = text.replace("\\%", "%").replace("\\", "")
    text = text.replace("Dem{v{s}}ar", "Demsar").replace('Ga{"e}l', "Gael")
    text = text.replace("Carr{'e}", "Carre").replace("Cram{'e}r", "Cramer's")
    text = text.replace("{", "").replace("}", "").replace("~", " ")
    text = re.sub(r"\s+([,.;:])", r"\1", text)
    return re.sub(r"\s+", " ", text).strip()


def add_figure(doc, filename, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(FIGURES / filename), width=Inches(6.2))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(6)
    for run in cp.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(8)
        run.italic = True


def main():
    tex = (PAPER / "ACP-Gini_Report_Springer.tex").read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Inches(8.5), Inches(11)
    section.top_margin = section.bottom_margin = Inches(.72)
    section.left_margin = section.right_margin = Inches(.78)
    section.header_distance = section.footer_distance = Inches(.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(10)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05
    for style_name, size in [("Heading 1", 12), ("Heading 2", 11), ("Heading 3", 10)]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(3)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("ACP-Gini: An Ancestor-Correlation Penalized Splitting Criterion for Multicollinearity-Robust Decision Trees")
    run.font.name = "Times New Roman"; run.font.size = Pt(16); run.bold = True
    authors = doc.add_paragraph(
        "Huynh Khang Lam | Tuan Dat Phan | Trong Nhan Nguyen | Xuan Huy Nguyen\n"
        "Artificial Intelligence, FPT University, Ho Chi Minh City, Vietnam\n"
        "khanglhse200666@fpt.edu.vn | ptundatwork@gmail.com | hi.gesar@gmail.com | HuyNX23@fe.edu.vn"
    )
    authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
    authors.paragraph_format.space_after = Pt(6)

    abstract_match = re.search(r"\\abstract\{(.*?)\}\s*\\keywords", tex, re.S)
    if abstract_match is None:
        abstract_match = re.search(r"\\begin\{abstract\}(.*?)\\keywords", tex, re.S)
    abstract = abstract_match.group(1)
    p = doc.add_paragraph()
    p.add_run("Abstract. ").bold = True
    p.add_run(clean_latex(abstract))
    keywords = re.search(r"\\keywords\{(.*?)\}", tex, re.S).group(1).replace("\\and", ",")
    p = doc.add_paragraph()
    p.add_run("Keywords: ").bold = True
    p.add_run(clean_latex(keywords))

    bibliography_start = tex.find("\\bibliographystyle")
    if bibliography_start < 0:
        bibliography_start = tex.index("\\bibliography{")
    body = tex[tex.index("\\section{Introduction}"):bibliography_start]
    body = re.sub(r"\\begin\{figure\*?\}.*?\\end\{figure\*?\}", "", body, flags=re.S)
    body = re.sub(r"\\begin\{table\*?\}.*?\\end\{table\*?\}", "", body, flags=re.S)
    body = re.sub(r"\\begin\{align\}.*?\\end\{align\}", "", body, flags=re.S)
    body = re.sub(r"\\begin\{equation\}.*?\\end\{equation\}", "", body, flags=re.S)
    body = body.replace("\\begin{itemize}", "").replace("\\end{itemize}", "")
    tokens = re.split(r"(\\section\{[^}]+\}|\\subsection\{[^}]+\}|\\paragraph\{[^}]+\}|\\item)", body)
    current = None
    for token in tokens:
        if not token.strip():
            continue
        match = re.match(r"\\(section|subsection|paragraph)\{([^}]+)\}", token)
        if match:
            level = {"section": 1, "subsection": 2, "paragraph": 3}[match.group(1)]
            doc.add_heading(match.group(2), level=level)
            current = match.group(1)
        elif token == "\\item":
            current = "item"
        else:
            for raw in re.split(r"\n\s*\n", token):
                cleaned = clean_latex(raw)
                if not cleaned:
                    continue
                if current == "item":
                    p = doc.add_paragraph(style="List Bullet")
                    p.add_run(cleaned)
                    current = None
                else:
                    doc.add_paragraph(cleaned)

    doc.add_heading("Generated Evidence Tables", level=1)
    positioning = pd.DataFrame([
        ["VIF+CART", "Yes", "Yes", "No", "No", "No"],
        ["RRF-style", "Yes", "No", "No", "Yes", "No"],
        ["ACP-Gini", "Yes", "Yes", "Yes", "Yes", "Yes"],
    ], columns=["Method", "Single tree", "Correlation-aware", "Path-local", "No preprocessing", "CART-compatible at alpha=0"])
    add_table(doc, "Related-work positioning.", positioning, [1.0, .75, 1.1, .8, 1.1, 1.45])
    stats = pd.read_csv(RESULTS / "dataset_stats.csv")
    add_table(doc, "Table 1. Data-set statistics.", stats.round(3), [1.3, .65, .55, 1.0, 1.8])
    main_frame = pd.read_csv(RESULTS / "uci_main.csv")
    pred = main_frame.groupby(["dataset", "method"])[["accuracy", "macro_f1", "auc"]].mean().reset_index().round(3)
    add_table(doc, "Table 2. Mean predictive performance over 15 folds.", pred, [1.0, 1.05, .8, .8, .8])
    redundancy = main_frame.groupby(["dataset", "method"])[["bootstrap_set_redundancy", "bootstrap_weighted_redundancy"]].mean().reset_index().round(3)
    add_table(doc, "Table 3. Bootstrap selected-set redundancy.", redundancy, [1.05, 1.15, 1.65, 1.85])
    stability = main_frame.groupby(["dataset", "method"])[["feature_set_jaccard", "top5_jaccard", "importance_rank_corr", "structural_distance"]].mean().reset_index().round(3)
    add_table(doc, "Table 4. Bootstrap stability metrics.", stability, [1.0, 1.0, 1.15, 1.0, 1.25, 1.15])

    doc.add_heading("Figures", level=1)
    add_figure(doc, "fig2_synthetic.png", "Figure 2. Synthetic group coverage and concentration.")
    add_figure(doc, "fig3_sensitivity.png", "Figure 3. Alpha sensitivity and noise importance.")
    add_figure(doc, "fig4_redundancy.png", "Figure 4. Bootstrap weighted redundancy.")
    add_figure(doc, "fig5_tree_comparison.png", "Figure 5. WDBC tree comparison.")
    add_figure(doc, "fig6_real_alpha_sweep.png", "Figure 6. WDBC and Wine operating characteristics.")

    doc.add_heading("References", level=1)
    bib = (PAPER / "references.bib").read_text(encoding="utf-8")
    entries = re.findall(r"@\w+\{[^,]+,(.*?)(?=\n\})", bib, re.S)
    for i, entry in enumerate(entries, 1):
        author = re.search(r"author=\{(.*?)\}", entry, re.S)
        title = re.search(r"title=\{(.*?)\}", entry, re.S)
        year = re.search(r"year=\{(.*?)\}", entry, re.S)
        text = f"[{i}] {author.group(1) if author else ''}. {title.group(1) if title else ''}. {year.group(1) if year else ''}."
        doc.add_paragraph(clean_latex(text))

    header = section.header.paragraphs[0]
    header.text = "ACP-Gini | FPT University"
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header.runs:
        run.font.name = "Times New Roman"; run.font.size = Pt(8); run.font.color.rgb = RGBColor(90, 90, 90)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field = OxmlElement("w:fldSimple"); field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)

    out = PAPER / "ACP-Gini_Report_Springer.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    main()
